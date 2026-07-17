import yaml
from pathlib import Path
from app.models.job import Job
import docx
from datetime import datetime
import re

class ResumeBuilder:
    def __init__(self):
        self.profile = self.load_profile()

    def load_profile(self):
        try:
            with open('config/profile.yaml', 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        except:
            return {'name': 'Aygun Aliyeva'}

    def sanitize_filename(self, text):
        text = re.sub(r'[\\/*?:"<>|]', '', text)
        return text.strip()[:60]

    def generate_for_job(self, job: Job, output_dir: str = 'outputs/resumes') -> str:
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        doc = docx.Document()
        doc.add_heading(self.profile.get('name', 'Aygun Aliyeva'), 0)
        doc.add_paragraph(f"{self.profile.get('email')} | {self.profile.get('phone')} | Remote")
        
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(f"Experienced Technical Writer and Content Specialist seeking the {job.title} role at {job.company}.")
        
        doc.add_heading('Core Skills', level=1)
        doc.add_paragraph('Technical Writing, Project Management, Content Creation, Research, Documentation, Python')
        
        safe_company = self.sanitize_filename(job.company)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{output_dir}/resume_{safe_company}_{timestamp}.docx'
        
        doc.save(filename)
        print(f'✅ Resume generated: {filename}')
        return filename
